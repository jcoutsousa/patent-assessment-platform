"""
Document Processing Pipeline for Patent Assessment Platform
Handles text extraction from various document formats
"""

import os
import io
import hashlib
import logging
from typing import Optional, Dict, Any, List
from datetime import datetime
from pathlib import Path

# Document processing libraries
import PyPDF2
from docx import Document as DocxDocument
from PIL import Image
import pytesseract
import aiofiles

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class DocumentProcessor:
    """
    Main document processing class that handles various file formats
    and extracts text content for patent analysis
    """

    def __init__(self, storage_path: str = "./uploads"):
        """
        Initialize document processor with storage configuration

        Args:
            storage_path: Directory path for temporary file storage
        """
        self.storage_path = Path(storage_path)
        self.storage_path.mkdir(exist_ok=True)

        # Supported file types and their processors
        self.processors = {
            'application/pdf': self.process_pdf,
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': self.process_docx,
            'text/plain': self.process_text,
            'image/png': self.process_image,
            'image/jpeg': self.process_image,
            'image/jpg': self.process_image,
        }

    async def process_document(self, file_content: bytes, filename: str, content_type: str) -> Dict[str, Any]:
        """
        Main entry point for document processing

        Args:
            file_content: Raw file bytes
            filename: Original filename
            content_type: MIME type of the file

        Returns:
            Dictionary containing extracted text and metadata
        """
        start_time = datetime.utcnow()

        # Calculate file hash for deduplication
        file_hash = hashlib.sha256(file_content).hexdigest()

        # Get the appropriate processor
        processor = self.processors.get(content_type)
        if not processor:
            raise ValueError(f"Unsupported file type: {content_type}")

        try:
            # Process the document
            result = await processor(file_content, filename)

            # Calculate processing time
            processing_time = (datetime.utcnow() - start_time).total_seconds()

            # Add metadata
            result.update({
                'filename': filename,
                'file_hash': file_hash,
                'file_size': len(file_content),
                'content_type': content_type,
                'processing_time_seconds': processing_time,
                'processed_at': datetime.utcnow().isoformat(),
                'status': 'success'
            })

            logger.info(f"Successfully processed {filename} in {processing_time:.2f}s")
            return result

        except Exception as e:
            logger.error(f"Error processing {filename}: {str(e)}")
            return {
                'filename': filename,
                'file_hash': file_hash,
                'status': 'error',
                'error': str(e),
                'processed_at': datetime.utcnow().isoformat()
            }

    async def process_pdf(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """
        Extract text from PDF documents using PyPDF2

        Args:
            file_content: PDF file bytes
            filename: Original filename

        Returns:
            Dictionary with extracted text and metadata
        """
        extracted_text = []
        metadata = {}
        page_count = 0
        has_images = False

        try:
            # Create PDF reader object
            pdf_file = io.BytesIO(file_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)

            # Extract metadata
            if pdf_reader.metadata:
                metadata = {
                    'title': pdf_reader.metadata.get('/Title', ''),
                    'author': pdf_reader.metadata.get('/Author', ''),
                    'subject': pdf_reader.metadata.get('/Subject', ''),
                    'creator': pdf_reader.metadata.get('/Creator', ''),
                    'creation_date': str(pdf_reader.metadata.get('/CreationDate', '')),
                }

            page_count = len(pdf_reader.pages)

            # Extract text from each page
            for page_num, page in enumerate(pdf_reader.pages, 1):
                try:
                    text = page.extract_text()
                    if text.strip():
                        extracted_text.append({
                            'page': page_num,
                            'text': text.strip()
                        })

                    # Check for images (for OCR fallback)
                    if '/XObject' in page['/Resources']:
                        x_objects = page['/Resources']['/XObject'].get_object()
                        for obj in x_objects:
                            if x_objects[obj]['/Subtype'] == '/Image':
                                has_images = True
                                break

                except Exception as e:
                    logger.warning(f"Error extracting page {page_num}: {str(e)}")
                    continue

            # Combine all text
            full_text = '\n\n'.join([p['text'] for p in extracted_text])

            # Calculate statistics
            word_count = len(full_text.split())
            char_count = len(full_text)

            return {
                'extracted_text': full_text,
                'page_texts': extracted_text,
                'metadata': metadata,
                'statistics': {
                    'page_count': page_count,
                    'word_count': word_count,
                    'character_count': char_count,
                    'has_images': has_images,
                    'extraction_method': 'PyPDF2'
                }
            }

        except Exception as e:
            logger.error(f"PDF processing error: {str(e)}")
            raise

    async def process_docx(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """
        Extract text from DOCX documents

        Args:
            file_content: DOCX file bytes
            filename: Original filename

        Returns:
            Dictionary with extracted text and metadata
        """
        try:
            # Create document object
            doc_file = io.BytesIO(file_content)
            doc = DocxDocument(doc_file)

            # Extract text from paragraphs
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text.strip())

            # Extract text from tables
            table_texts = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        table_texts.append(' | '.join(row_text))

            # Combine all text
            full_text = '\n\n'.join(paragraphs)
            if table_texts:
                full_text += '\n\nTables:\n' + '\n'.join(table_texts)

            # Extract metadata
            core_props = doc.core_properties
            metadata = {
                'title': core_props.title or '',
                'author': core_props.author or '',
                'subject': core_props.subject or '',
                'created': str(core_props.created) if core_props.created else '',
                'modified': str(core_props.modified) if core_props.modified else '',
            }

            # Calculate statistics
            word_count = len(full_text.split())
            char_count = len(full_text)

            return {
                'extracted_text': full_text,
                'paragraphs': paragraphs,
                'tables': table_texts,
                'metadata': metadata,
                'statistics': {
                    'paragraph_count': len(paragraphs),
                    'table_count': len(doc.tables),
                    'word_count': word_count,
                    'character_count': char_count,
                    'extraction_method': 'python-docx'
                }
            }

        except Exception as e:
            logger.error(f"DOCX processing error: {str(e)}")
            raise

    async def process_text(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """
        Process plain text files

        Args:
            file_content: Text file bytes
            filename: Original filename

        Returns:
            Dictionary with text content
        """
        try:
            # Decode text with multiple encoding attempts
            encodings = ['utf-8', 'latin-1', 'ascii', 'utf-16']
            text = None

            for encoding in encodings:
                try:
                    text = file_content.decode(encoding)
                    break
                except UnicodeDecodeError:
                    continue

            if text is None:
                raise ValueError("Unable to decode text file")

            # Calculate statistics
            lines = text.split('\n')
            word_count = len(text.split())
            char_count = len(text)

            return {
                'extracted_text': text,
                'statistics': {
                    'line_count': len(lines),
                    'word_count': word_count,
                    'character_count': char_count,
                    'extraction_method': 'direct'
                }
            }

        except Exception as e:
            logger.error(f"Text processing error: {str(e)}")
            raise

    async def process_image(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """
        Extract text from images using OCR (Tesseract)

        Args:
            file_content: Image file bytes
            filename: Original filename

        Returns:
            Dictionary with extracted text from OCR
        """
        try:
            # Open image
            image = Image.open(io.BytesIO(file_content))

            # Convert to RGB if necessary
            if image.mode != 'RGB':
                image = image.convert('RGB')

            # Perform OCR
            extracted_text = pytesseract.image_to_string(image)

            # Get image metadata
            width, height = image.size

            # Calculate statistics
            word_count = len(extracted_text.split())
            char_count = len(extracted_text)

            return {
                'extracted_text': extracted_text.strip(),
                'metadata': {
                    'dimensions': f"{width}x{height}",
                    'mode': image.mode,
                    'format': image.format
                },
                'statistics': {
                    'word_count': word_count,
                    'character_count': char_count,
                    'extraction_method': 'OCR (Tesseract)'
                }
            }

        except Exception as e:
            logger.error(f"Image OCR error: {str(e)}")
            raise

    async def preprocess_text(self, text: str) -> Dict[str, Any]:
        """
        Preprocess extracted text for patent analysis

        Args:
            text: Raw extracted text

        Returns:
            Dictionary with preprocessed text and features
        """
        # Remove excessive whitespace
        cleaned_text = ' '.join(text.split())

        # Extract potential technical terms (simple heuristic)
        words = cleaned_text.lower().split()
        technical_terms = [w for w in words if len(w) > 10 or '-' in w or '_' in w]

        # Identify potential sections
        sections = {}
        section_keywords = ['abstract', 'background', 'summary', 'description',
                          'claims', 'embodiment', 'invention', 'technical field']

        for keyword in section_keywords:
            if keyword in text.lower():
                # Find the section (simplified)
                start_idx = text.lower().find(keyword)
                if start_idx != -1:
                    sections[keyword] = True

        return {
            'cleaned_text': cleaned_text,
            'features': {
                'technical_terms': list(set(technical_terms[:50])),  # Top 50 unique terms
                'identified_sections': list(sections.keys()),
                'text_length': len(cleaned_text),
                'sentence_count': cleaned_text.count('.') + cleaned_text.count('!') + cleaned_text.count('?')
            }
        }


# Async file handling utilities
async def save_uploaded_file(file_content: bytes, filename: str, storage_path: str = "./uploads") -> str:
    """
    Save uploaded file to disk

    Args:
        file_content: File bytes
        filename: Original filename
        storage_path: Directory to save files

    Returns:
        Path to saved file
    """
    storage_dir = Path(storage_path)
    storage_dir.mkdir(exist_ok=True)

    # Generate unique filename
    timestamp = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
    file_hash = hashlib.sha256(file_content).hexdigest()[:8]
    safe_filename = f"{timestamp}_{file_hash}_{filename}"
    file_path = storage_dir / safe_filename

    # Save file
    async with aiofiles.open(file_path, 'wb') as f:
        await f.write(file_content)

    return str(file_path)


async def delete_uploaded_file(file_path: str) -> bool:
    """
    Delete uploaded file after processing

    Args:
        file_path: Path to file

    Returns:
        True if deleted successfully
    """
    try:
        if os.path.exists(file_path):
            os.remove(file_path)
            return True
    except Exception as e:
        logger.error(f"Error deleting file {file_path}: {str(e)}")
    return False